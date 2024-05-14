
import tkinter as tk 

from tkinter  import *
import docx
from docx import Document
from docx.shared import Pt
import pandas as pd
from tkinter import ttk
from datetime import date

today = date.today()

# Top level window 
frame = tk.Tk() 
frame.title("TextBox Input") 
frame.geometry('800x400') 
# Function for getting Input 
# from textbox and printing it  
# at label widget 

currentIndex = ""
tansactionData = {
    "Date" : [],
    "Acitivity" : [],
    "Amount" : []
}
accountingEntriesData = {
    "Date" : [],
    "Activity" : [],
    "Account Name" : [],
    "Debit" : [],
    "Credit": [],
    "Nature of Account" : [],
}

def OptionCallBack(*args):
    global currentIndex
    currentIndex = variable.get()


def handleSubmit(): 
    global currentIndex, tansactionData, accountingEntriesData
    name = inputtxt.get(1.0, "end-1c") 
    value = 8000 if currentIndex == "Individual" else 12000
    tansactionData["Date"].append(today)
    tansactionData["Acitivity"].append( "Account Opening Fee- Customer " + name)
    tansactionData["Amount"].append(f'{value:,}'    )

    accountingEntriesData["Date"].append(today)
    accountingEntriesData["Activity"].append("Account Opening Fee- Customer " + name)
    accountingEntriesData["Account Name"].append("Customer Account - Customer " + name + " - USD")
    accountingEntriesData["Debit"].append(f'{value:,}'    )
    accountingEntriesData["Credit"].append("")
    accountingEntriesData["Nature of Account"].append("Customer Account - USD")

    accountingEntriesData["Date"].append("")
    accountingEntriesData["Activity"].append("")
    accountingEntriesData["Account Name"].append("Onboarding Fee - " + currentIndex)
    accountingEntriesData["Debit"].append("")
    accountingEntriesData["Credit"].append(f'{value:,}'    )
    accountingEntriesData["Nature of Account"].append("Revenue")



def handleExport():

    global tansactionData, accountingEntriesData


    df = pd.DataFrame(tansactionData)
    print(df)

    doc = docx.Document()
    doc.add_heading('Transaction')

    t = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])
    # Add borders
    t.style = 'TableGrid'

    # Add the column headings
    for j in range(df.shape[1]):
        t.cell(0, j).text = df.columns[j]

    # Add the body of the data frame to the table
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            cell = df.iat[i, j]
            t.cell(i+1, j).text = str(cell)

    doc.add_heading('Accounting Entries')
    df2 = pd.DataFrame(accountingEntriesData)
    table2 =  doc.add_table(rows=df2.shape[0]+1, cols=df2.shape[1])
     # Add borders
    table2.style = 'TableGrid'

    # Add the column headings
    for j in range(df2.shape[1]):
        table2.cell(0, j).text = df2.columns[j]

    # Add the body of the data frame to the table
    for i in range(df2.shape[0]):
        for j in range(df2.shape[1]):
            cell = df2.iat[i, j]
            table2.cell(i+1, j).text = str(cell)
    
    doc.save('table 1.docx')

   
    


user_name = Label(frame, 
                  text = "Customer Name").place(x = 30,
                                           y = 60)  

# TextBox Creation 
inputtxt = tk.Text(frame, 
                   height = 1, 
                   width = 18)
  
inputtxt.pack() 
inputtxt.place(x = 150, y = 60)


user_type = Label(frame, 
                  text = "Customer Type").place(x = 30,
                                           y = 100)  




variable = StringVar(frame)
variable.set("Select From List")
variable.trace('w', OptionCallBack)



types = ["Individual", "Corporate"]
combobox = ttk.Combobox(values=types, textvariable=variable,)
combobox.pack(anchor=NW, padx=6, pady=6)
combobox.place(
    x = 150,
    y = 100
)


  
# Button Creation 
printButton = tk.Button(frame, 
                        text = "Submit",  
                        command = handleSubmit).place(x=120, y = 150) 

exportButton = tk.Button(frame, 
                        text = "Export to Doc",  
                        command = handleExport).place(x=180, y = 150) 

frame.mainloop() 